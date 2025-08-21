"""
文档解析智能体
负责解析Word、PDF等需求文档，提取测试用例相关信息
"""
import uuid
import json
import asyncio
from typing import Dict, List, Any, Optional
from datetime import datetime
from pathlib import Path

from autogen_agentchat.base import TaskResult
from autogen_agentchat.messages import ModelClientStreamingChunkEvent
from autogen_core import message_handler, type_subscription, MessageContext, TopicId
from loguru import logger
from pydantic import BaseModel, Field

from app.core.agents.base import BaseAgent
from app.core.types import TopicTypes, AgentTypes, AGENT_NAMES
from app.core.messages.test_case import (
    DocumentParseRequest, DocumentParseResponse,
    TestCaseData
)
from app.core.enums import TestType, TestLevel, Priority, InputSource
from app.agents.database.requirement_saver_agent import RequirementSaveRequest


class DocumentParseResult(BaseModel):
    """文档解析结果"""
    document_type: str = Field(..., description="文档类型")
    title: str = Field(..., description="文档标题")
    content: str = Field(..., description="文档内容")
    sections: List[Dict[str, Any]] = Field(default_factory=list, description="文档章节")
    requirements: List[Dict[str, Any]] = Field(default_factory=list, description="需求列表")
    test_scenarios: List[Dict[str, Any]] = Field(default_factory=list, description="测试场景")
    confidence_score: float = Field(0.0, description="解析置信度")


@type_subscription(topic_type=TopicTypes.DOCUMENT_PARSER.value)
class DocumentParserAgent(BaseAgent):
    """文档解析智能体，负责解析各种格式的需求文档"""

    def __init__(self, model_client_instance=None, **kwargs):
        """初始化文档解析智能体"""
        super().__init__(
            agent_id=AgentTypes.DOCUMENT_PARSER.value,
            agent_name=AGENT_NAMES.get(AgentTypes.DOCUMENT_PARSER.value, "文档解析智能体"),
            model_client_instance=model_client_instance,
            **kwargs
        )
        
        # 支持的文档格式
        self.supported_formats = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.doc': self._parse_doc,
            '.txt': self._parse_txt,
            '.md': self._parse_markdown
        }
        
        logger.info(f"文档解析智能体初始化完成: {self.agent_name}")

    @message_handler
    async def handle_document_parse_request(self, message: DocumentParseRequest, ctx: MessageContext) -> None:
        """处理文档解析请求"""
        start_time = datetime.now()

        try:
            logger.info(f"开始处理文档解析请求: {message.session_id}")

            # 发送开始处理消息
            await self.send_response(
                f"🔍 开始解析文档: {message.file_name}",
                region="process"
            )

            # 发送文档信息
            file_path = Path(message.file_path)
            file_size = file_path.stat().st_size if file_path.exists() else 0
            await self.send_response(
                f"📄 文档信息: 文件大小 {file_size / 1024:.1f}KB, 格式: {file_path.suffix}",
                region="info"
            )

            # 解析文档
            await self.send_response("🔄 第1步: 开始解析文档内容...", region="progress")
            parse_result = await self._parse_document(message)

            # 发送解析结果统计
            await self.send_response(
                f"📊 解析结果: 提取 {len(parse_result.sections)} 个章节, {len(parse_result.requirements)} 个需求",
                region="info",
                result={
                    "sections_count": len(parse_result.sections),
                    "requirements_count": len(parse_result.requirements),
                    "confidence_score": parse_result.confidence_score
                }
            )

            # 生成测试用例
            await self.send_response("🔄 第2步: 基于文档内容生成测试用例...", region="progress")
            test_cases = await self._generate_test_cases_from_document(
                parse_result, message
            )

            # 发送测试用例生成结果
            await self.send_response(
                f"✅ 成功生成 {len(test_cases)} 个测试用例",
                region="success",
                result={"test_cases_count": len(test_cases)}
            )

            # 保存需求到数据库
            if parse_result.requirements:
                await self.send_response("🔄 第3步: 保存需求信息到数据库...", region="progress")
                await self._save_requirements_to_database(parse_result, message)

            # 计算处理时间
            processing_time = (datetime.now() - start_time).total_seconds()

            # 构建响应
            response = DocumentParseResponse(
                session_id=message.session_id,
                document_id=str(uuid.uuid4()),
                file_name=message.file_name,
                file_path=message.file_path,
                parse_result=parse_result.model_dump(),
                test_cases=test_cases,
                processing_time=processing_time,
                created_at=datetime.now().isoformat()
            )

            # 发送完成消息
            await self.send_response(
                f"✅ 文档解析完成! 处理时间: {processing_time:.2f}秒",
                is_final=False,
                region="success",
                result={
                    "processing_time": processing_time,
                    "total_test_cases": len(test_cases),
                    "total_requirements": len(parse_result.requirements),
                    "confidence_score": parse_result.confidence_score
                }
            )

            # 发送到测试点提取智能体
            await self.send_response("🔄 转发到测试点提取智能体进行专业测试点分析...", region="info")
            await self._send_to_test_point_extractor(response)

        except Exception as e:
            processing_time = (datetime.now() - start_time).total_seconds()
            logger.error(f"文档解析失败: {str(e)}")
            await self.send_response(
                f"❌ 文档解析失败: {str(e)} (处理时间: {processing_time:.2f}秒)",
                is_final=True,
                region="error",
                result={"processing_time": processing_time, "error": str(e)}
            )

    async def _parse_document(self, message: DocumentParseRequest) -> DocumentParseResult:
        """解析文档内容"""
        try:
            file_path = Path(message.file_path)
            file_extension = file_path.suffix.lower()

            if file_extension not in self.supported_formats:
                await self.send_response(
                    f"❌ 不支持的文档格式: {file_extension}",
                    region="error"
                )
                raise ValueError(f"不支持的文档格式: {file_extension}")

            await self.send_response(
                f"📖 正在解析 {file_extension} 格式文档...",
                region="progress"
            )

            # 调用对应的解析方法
            parser_func = self.supported_formats[file_extension]
            content_start_time = datetime.now()

            # 提取页面内容
            content = await parser_func(file_path)
            content_time = (datetime.now() - content_start_time).total_seconds()

            await self.send_response(
                f"📝 内容提取完成，耗时 {content_time:.2f}秒，内容长度: {len(content)} 字符",
                region="info",
                result={"content_extraction_time": content_time, "content_length": len(content)}
            )

            # 使用AI分析文档内容
            await self.send_response("🤖 开始AI智能分析文档内容...", region="progress")
            ai_start_time = datetime.now()
            analysis_result = await self._analyze_document_content(content, message)
            ai_time = (datetime.now() - ai_start_time).total_seconds()

            await self.send_response(
                f"🧠 AI分析完成，耗时 {ai_time:.2f}秒，置信度: {analysis_result.confidence_score:.2f}",
                region="success",
                result={
                    "ai_analysis_time": ai_time,
                    "confidence_score": analysis_result.confidence_score,
                    "analysis_quality": "高" if analysis_result.confidence_score > 0.8 else "中" if analysis_result.confidence_score > 0.6 else "低"
                }
            )

            return analysis_result

        except Exception as e:
            logger.error(f"文档解析失败: {str(e)}")
            await self.send_response(
                f"❌ 文档解析过程中发生错误: {str(e)}",
                region="error"
            )
            raise
    async def _parse_pdf(self, file_path: Path) -> str:
        """解析PDF文档，提取完整文本内容"""
        try:
            # 导入必要的库
            import fitz  # PyMuPDF

            logger.info(f"开始解析PDF文档: {file_path}")

            # 打开PDF文档
            pdf_document = fitz.open(file_path)
            total_pages = len(pdf_document)

            await self.send_response(
                f"📄 PDF文档信息: 共 {total_pages} 页，开始文本提取...",
                region="info",
                result={"total_pages": total_pages, "parsing_method": "text_extraction"}
            )

            # 存储所有页面的文本内容
            all_text = []
            successful_pages = 0
            failed_pages = 0
            total_content_length = 0

            # 逐页提取文本
            for page_num in range(total_pages):
                page_start_time = datetime.now()

                try:
                    await self.send_response(
                        f"🔍 正在提取第 {page_num + 1}/{total_pages} 页文本...",
                        region="progress"
                    )

                    # 获取页面
                    page = pdf_document[page_num]

                    # 提取文本内容
                    page_text = page.get_text()

                    # 清理和格式化文本
                    cleaned_text = self._clean_pdf_text(page_text)

                    page_time = (datetime.now() - page_start_time).total_seconds()

                    if cleaned_text.strip():
                        all_text.append({
                            'page_number': page_num + 1,
                            'content': cleaned_text.strip()
                        })
                        successful_pages += 1
                        total_content_length += len(cleaned_text)

                        await self.send_response(
                            f"✅ 第 {page_num + 1} 页文本提取成功 (耗时: {page_time:.2f}秒, 内容: {len(cleaned_text)} 字符)",
                            region="success"
                        )
                    else:
                        failed_pages += 1
                        await self.send_response(
                            f"⚠️ 第 {page_num + 1} 页无文本内容 (耗时: {page_time:.2f}秒)",
                            region="warning"
                        )

                    logger.info(f"第 {page_num + 1} 页文本提取完成，内容长度: {len(cleaned_text)}")

                except Exception as e:
                    failed_pages += 1
                    page_time = (datetime.now() - page_start_time).total_seconds()
                    logger.error(f"第 {page_num + 1} 页文本提取失败: {str(e)}")
                    await self.send_response(
                        f"❌ 第 {page_num + 1} 页文本提取失败: {str(e)} (耗时: {page_time:.2f}秒)",
                        region="error"
                    )
                    # 继续处理下一页，不中断整个流程
                    continue

            # 关闭PDF文档
            pdf_document.close()

            # 发送提取统计
            await self.send_response(
                f"📊 PDF文本提取统计: 成功 {successful_pages} 页, 失败 {failed_pages} 页, 总内容 {total_content_length} 字符",
                region="info",
                result={
                    "successful_pages": successful_pages,
                    "failed_pages": failed_pages,
                    "total_content_length": total_content_length,
                    "success_rate": (successful_pages / total_pages * 100) if total_pages > 0 else 0
                }
            )

            # 整合所有页面内容
            await self.send_response("🔄 正在整合页面内容...", region="progress")
            integration_start_time = datetime.now()
            final_content = self._integrate_pdf_text_content(all_text, file_path.name)
            integration_time = (datetime.now() - integration_start_time).total_seconds()

            await self.send_response(
                f"✅ PDF文本解析完成! 有效页面: {len(all_text)}, 内容整合耗时: {integration_time:.2f}秒",
                region="success",
                result={
                    "effective_pages": len(all_text),
                    "integration_time": integration_time,
                    "final_content_length": len(final_content)
                }
            )

            return final_content

        except ImportError as e:
            if "fitz" in str(e):
                raise ImportError("需要安装 PyMuPDF: pip install PyMuPDF")
            else:
                raise ImportError(f"缺少必要依赖: {str(e)}")
        except Exception as e:
            logger.error(f"PDF解析失败: {str(e)}")
            raise

    def _clean_pdf_text(self, text: str) -> str:
        """清理和格式化PDF提取的文本"""
        try:
            if not text:
                return ""

            # 移除多余的空白字符
            lines = text.split('\n')
            cleaned_lines = []

            for line in lines:
                # 移除行首行尾空白
                line = line.strip()

                # 跳过空行
                if not line:
                    continue

                # 移除多余的空格
                line = ' '.join(line.split())

                # 跳过过短的行（可能是页眉页脚或噪音）
                if len(line) < 3:
                    continue

                cleaned_lines.append(line)

            # 合并相关行（处理PDF中的换行问题）
            merged_lines = []
            current_paragraph = ""

            for line in cleaned_lines:
                # 如果行以句号、问号、感叹号等结尾，认为是段落结束
                if line.endswith(('.', '。', '!', '！', '?', '？', ':', '：')):
                    current_paragraph += " " + line if current_paragraph else line
                    merged_lines.append(current_paragraph.strip())
                    current_paragraph = ""
                # 如果行看起来像标题（短且可能包含数字）
                elif len(line) < 50 and any(char.isdigit() for char in line[:10]):
                    if current_paragraph:
                        merged_lines.append(current_paragraph.strip())
                        current_paragraph = ""
                    merged_lines.append(line)
                else:
                    current_paragraph += " " + line if current_paragraph else line

            # 添加最后的段落
            if current_paragraph:
                merged_lines.append(current_paragraph.strip())

            return '\n\n'.join(merged_lines)

        except Exception as e:
            logger.error(f"清理PDF文本失败: {str(e)}")
            return text  # 返回原始文本

    def _integrate_pdf_text_content(self, page_contents: list, filename: str) -> str:
        """整合所有页面的文本内容为完整文档"""
        try:
            # 构建文档头部
            content_lines = [
                f"# {filename} - 文档内容",
                "",
                f"*本文档由PDF文本提取自动生成，共{len(page_contents)}页*",
                "",
                "---",
                ""
            ]

            # 如果页面较多，添加简单目录
            if len(page_contents) > 5:
                content_lines.extend([
                    "## 目录",
                    ""
                ])
                for page_info in page_contents:
                    page_num = page_info['page_number']
                    # 尝试从内容中提取前几个词作为目录项
                    content_preview = page_info['content'][:100].replace('\n', ' ')
                    if len(content_preview) > 50:
                        content_preview = content_preview[:50] + "..."
                    content_lines.append(f"- 第{page_num}页: {content_preview}")

                content_lines.extend(["", "---", ""])

            # 添加每页内容
            for page_info in page_contents:
                page_num = page_info['page_number']
                content = page_info['content']

                # 添加页面标题
                content_lines.extend([
                    f"## 第{page_num}页",
                    ""
                ])

                # 添加页面内容
                content_lines.extend([
                    content,
                    "",
                    "---",
                    ""
                ])

            # 添加文档尾部
            content_lines.extend([
                "",
                f"*文档解析完成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*"
            ])

            return "\n".join(content_lines)

        except Exception as e:
            logger.error(f"整合PDF文本内容失败: {str(e)}")
            # 返回简单的内容拼接
            simple_content = f"# {filename}\n\n"
            for page_info in page_contents:
                simple_content += f"## 第{page_info['page_number']}页\n\n{page_info['content']}\n\n---\n\n"
            return simple_content

    async def _parse_pdf_with_multimodal(self, file_path: Path) -> str:
        """使用多模态大模型解析PDF文档"""
        try:
            # 导入必要的库
            import fitz  # PyMuPDF
            from io import BytesIO
            from autogen_agentchat.messages import MultiModalMessage
            from autogen_core import Image as AGImage
            from PIL import Image

            logger.info(f"开始使用多模态模型解析PDF: {file_path}")

            # 打开PDF文档
            pdf_document = fitz.open(file_path)
            total_pages = len(pdf_document)

            await self.send_response(
                f"📄 PDF文档信息: 共 {total_pages} 页，开始多模态解析...",
                region="info",
                result={"total_pages": total_pages, "parsing_method": "multimodal"}
            )

            # 存储每页解析结果
            page_contents = []
            successful_pages = 0
            failed_pages = 0
            total_content_length = 0

            # 逐页处理
            for page_num in range(total_pages):
                page_start_time = datetime.now()

                try:
                    await self.send_response(
                        f"🔍 正在解析第 {page_num + 1}/{total_pages} 页...",
                        region="progress"
                    )

                    # 获取页面
                    page = pdf_document[page_num]

                    # 将页面转换为图片
                    pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))  # 2倍缩放提高清晰度
                    img_data = pix.tobytes("png")
                    img_size = len(img_data)

                    # 转换为PIL图片
                    pil_image = Image.open(BytesIO(img_data))

                    # 转换为AGImage
                    ag_image = AGImage(pil_image)

                    # 使用多模态模型解析页面内容
                    page_content = await self._analyze_pdf_page_with_multimodal(
                        ag_image, page_num + 1, total_pages
                    )

                    page_time = (datetime.now() - page_start_time).total_seconds()

                    if page_content.strip():
                        page_contents.append({
                            'page_number': page_num + 1,
                            'content': page_content.strip()
                        })
                        successful_pages += 1
                        total_content_length += len(page_content)

                        await self.send_response(
                            f"✅ 第 {page_num + 1} 页解析成功 (耗时: {page_time:.2f}秒, 内容: {len(page_content)} 字符, 图片: {img_size/1024:.1f}KB)",
                            region="success"
                        )
                    else:
                        failed_pages += 1
                        await self.send_response(
                            f"⚠️ 第 {page_num + 1} 页内容为空 (耗时: {page_time:.2f}秒)",
                            region="warning"
                        )

                    logger.info(f"第 {page_num + 1} 页解析完成，内容长度: {len(page_content)}")

                except Exception as e:
                    failed_pages += 1
                    page_time = (datetime.now() - page_start_time).total_seconds()
                    logger.error(f"第 {page_num + 1} 页解析失败: {str(e)}")
                    await self.send_response(
                        f"❌ 第 {page_num + 1} 页解析失败: {str(e)} (耗时: {page_time:.2f}秒)",
                        region="error"
                    )
                    # 继续处理下一页，不中断整个流程
                    continue

            # 关闭PDF文档
            pdf_document.close()

            # 发送解析统计
            await self.send_response(
                f"📊 PDF解析统计: 成功 {successful_pages} 页, 失败 {failed_pages} 页, 总内容 {total_content_length} 字符",
                region="info",
                result={
                    "successful_pages": successful_pages,
                    "failed_pages": failed_pages,
                    "total_content_length": total_content_length,
                    "success_rate": (successful_pages / total_pages * 100) if total_pages > 0 else 0
                }
            )

            # 整合所有页面内容为markdown格式
            await self.send_response("🔄 正在整合页面内容为Markdown格式...", region="progress")
            integration_start_time = datetime.now()
            markdown_content = await self._integrate_pdf_pages_to_markdown(page_contents, file_path.name)
            integration_time = (datetime.now() - integration_start_time).total_seconds()

            await self.send_response(
                f"✅ PDF解析完成! 有效页面: {len(page_contents)}, 内容整合耗时: {integration_time:.2f}秒",
                region="success",
                result={
                    "effective_pages": len(page_contents),
                    "integration_time": integration_time,
                    "final_content_length": len(markdown_content)
                }
            )

            return markdown_content

        except ImportError as e:
            if "fitz" in str(e):
                raise ImportError("需要安装 PyMuPDF: pip install PyMuPDF")
            else:
                raise ImportError(f"缺少必要依赖: {str(e)}")
        except Exception as e:
            logger.error(f"PDF解析失败: {str(e)}")
            raise

    async def _analyze_pdf_page_with_multimodal(self, ag_image, page_num: int, total_pages: int) -> str:
        """使用多模态大模型分析PDF页面"""
        try:
            from autogen_agentchat.messages import MultiModalMessage
            from app.agents.factory import agent_factory

            # 创建多模态分析智能体（使用千问模型）
            multimodal_agent = agent_factory.create_assistant_agent(
                name="pdf_page_analyzer",
                system_message=self._build_pdf_analysis_system_prompt(),
                model_client_type="qwenvl"  # 使用千问VL模型
            )

            # 构建多模态消息
            analysis_prompt = f"""
请仔细分析这张PDF页面图片（第{page_num}页，共{total_pages}页），提取其中的文本内容和结构信息。

要求：
1. 识别并提取所有可见的文字内容
2. 保持原有的层次结构和格式
3. 识别标题、段落、列表、表格等元素
4. 如果有图表或图像，请描述其内容和作用
5. 输出格式为清晰的文本，保持逻辑结构

请直接输出页面内容，不需要额外的解释说明。
"""

            # 创建多模态消息
            multi_modal_message = MultiModalMessage(
                content=[analysis_prompt, ag_image],
                source="user"
            )

            # 执行分析
            result = await multimodal_agent.run(task=multi_modal_message)

            # 提取结果内容
            if result.messages and hasattr(result.messages[-1], 'content'):
                return result.messages[-1].content
            else:
                logger.warning(f"第{page_num}页多模态分析未返回有效内容")
                return ""

        except Exception as e:
            logger.error(f"第{page_num}页多模态分析失败: {str(e)}")
            return f"第{page_num}页解析失败: {str(e)}"

    def _build_pdf_analysis_system_prompt(self) -> str:
        """构建PDF页面分析系统提示"""
        return """
你是专业的文档内容识别专家，擅长从PDF页面图片中准确提取文本内容和结构信息。

你的任务是：
1. 仔细观察PDF页面图片
2. 识别并提取所有文字内容
3. 保持原有的文档结构和格式
4. 识别标题、段落、列表、表格等不同类型的内容
5. 对于图表、图像等非文字内容，从多个角度详细描述

输出要求：
- 直接输出页面的文本内容
- 保持清晰的层次结构
- 使用适当的标记来区分不同类型的内容
- 确保内容的完整性和准确性
- 不要添加多余的解释或评论
"""

    async def _integrate_pdf_pages_to_markdown(self, page_contents: list, filename: str) -> str:
        """将所有页面内容整合为完整的markdown文档"""
        try:
            # 构建markdown文档头部
            markdown_lines = [
                f"# {filename} - 文档内容解析",
                "",
                f"*本文档由AI多模态模型自动解析生成，共{len(page_contents)}页*",
                "",
                "---",
                ""
            ]

            # 添加目录（如果页面较多）
            if len(page_contents) > 3:
                markdown_lines.extend([
                    "## 目录",
                    ""
                ])
                for page_info in page_contents:
                    page_num = page_info['page_number']
                    # 尝试从内容中提取标题作为目录项
                    content_lines = page_info['content'].split('\n')
                    title = f"第{page_num}页"
                    for line in content_lines[:5]:  # 检查前5行寻找可能的标题
                        line = line.strip()
                        if line and len(line) < 100:  # 可能是标题
                            title = f"第{page_num}页: {line[:50]}..."
                            break
                    markdown_lines.append(f"- [{title}](#第{page_num}页)")

                markdown_lines.extend(["", "---", ""])

            # 添加每页内容
            for page_info in page_contents:
                page_num = page_info['page_number']
                content = page_info['content']

                # 添加页面标题
                markdown_lines.extend([
                    f"## 第{page_num}页",
                    ""
                ])

                # 处理页面内容，确保markdown格式正确
                processed_content = self._process_page_content_for_markdown(content)
                markdown_lines.extend([
                    processed_content,
                    "",
                    "---",
                    ""
                ])

            # 添加文档尾部
            markdown_lines.extend([
                "",
                f"*文档解析完成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*"
            ])

            return "\n".join(markdown_lines)

        except Exception as e:
            logger.error(f"整合PDF页面内容失败: {str(e)}")
            # 返回简单的内容拼接
            simple_content = f"# {filename}\n\n"
            for page_info in page_contents:
                simple_content += f"## 第{page_info['page_number']}页\n\n{page_info['content']}\n\n---\n\n"
            return simple_content

    def _process_page_content_for_markdown(self, content: str) -> str:
        """处理页面内容，确保markdown格式正确"""
        try:
            lines = content.split('\n')
            processed_lines = []

            for line in lines:
                line = line.strip()
                if not line:
                    processed_lines.append("")
                    continue

                # 处理可能的标题（以数字或特殊字符开头的行）
                if any(line.startswith(prefix) for prefix in ['1.', '2.', '3.', '4.', '5.', '一、', '二、', '三、', '（一）', '（二）']):
                    processed_lines.append(f"### {line}")
                # 处理列表项
                elif line.startswith(('•', '·', '-', '*')):
                    processed_lines.append(f"- {line[1:].strip()}")
                # 处理普通内容
                else:
                    processed_lines.append(line)

            return "\n".join(processed_lines)

        except Exception as e:
            logger.error(f"处理页面内容格式失败: {str(e)}")
            return content

    async def _parse_docx(self, file_path: Path) -> str:
        """解析DOCX文档"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            content = ""
            
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # 解析表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join([cell.text for cell in row.cells])
                    content += row_text + "\n"
            
            return content.strip()
            
        except ImportError:
            raise ImportError("需要安装 python-docx: pip install python-docx")
        except Exception as e:
            logger.error(f"DOCX解析失败: {str(e)}")
            raise

    async def _parse_doc(self, file_path: Path) -> str:
        """解析DOC文档"""
        try:
            import subprocess
            import tempfile
            
            # 使用LibreOffice转换为文本
            with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as tmp_file:
                cmd = [
                    'libreoffice', '--headless', '--convert-to', 'txt',
                    '--outdir', str(Path(tmp_file.name).parent),
                    str(file_path)
                ]
                
                subprocess.run(cmd, check=True, capture_output=True)
                
                txt_file = Path(tmp_file.name).with_suffix('.txt')
                if txt_file.exists():
                    content = txt_file.read_text(encoding='utf-8')
                    txt_file.unlink()  # 删除临时文件
                    return content
                else:
                    raise Exception("DOC转换失败")
                    
        except Exception as e:
            logger.error(f"DOC解析失败: {str(e)}")
            # 回退到简单文本读取
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            except:
                raise Exception(f"DOC文档解析失败: {str(e)}")

    async def _parse_txt(self, file_path: Path) -> str:
        """解析TXT文档"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            for encoding in ['gbk', 'gb2312', 'latin-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            raise Exception("无法识别文档编码")

    async def _parse_markdown(self, file_path: Path) -> str:
        """解析Markdown文档"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Markdown解析失败: {str(e)}")
            raise

    async def _analyze_document_content(self, content: str, message: DocumentParseRequest) -> DocumentParseResult:
        """使用AI分析文档内容"""
        try:
            # 创建AI分析智能体
            agent = self._create_document_analysis_agent()
            
            # 构建分析提示
            analysis_prompt = self._build_document_analysis_prompt(content, message)
            
            # 执行AI分析
            analysis_result = await self._run_ai_analysis(agent, analysis_prompt)
            
            # 解析AI响应
            return self._parse_ai_analysis_result(analysis_result, content)
            
        except Exception as e:
            logger.error(f"AI文档分析失败: {str(e)}")
            # 返回基础解析结果
            return DocumentParseResult(
                document_type="unknown",
                title=message.file_name,
                content=content,
                sections=[],
                requirements=[],
                test_scenarios=[],
                confidence_score=0.5
            )

    def _create_document_analysis_agent(self):
        """创建文档分析智能体"""
        from app.agents.factory import agent_factory

        return agent_factory.create_assistant_agent(
            name="document_analyzer",
            system_message=self._build_document_analysis_system_prompt(),
            model_client_type="deepseek"
        )

    def _build_document_analysis_system_prompt(self) -> str:
        """构建文档分析系统提示"""
        return """
你是专业的需求文档分析专家，擅长从各种文档中提取测试相关信息。

你的任务是：
1. 分析文档结构和内容
2. 识别功能需求和非功能需求
3. 提取测试场景和用例
4. 生成结构化的分析结果

请按照以下JSON格式返回分析结果：
{
    "document_type": "需求文档类型",
    "title": "文档标题",
    "sections": [
        {
            "title": "章节标题",
            "content": "章节内容摘要",
            "level": 1
        }
    ],
    "requirements": [
        {
            "id": "需求ID",
            "title": "需求标题",
            "description": "需求描述",
            "type": "functional/non-functional",
            "priority": "high/medium/low"
        }
    ],
    "test_scenarios": [
        {
            "title": "测试场景标题",
            "description": "测试场景描述",
            "requirements": ["关联需求ID"],
            "test_type": "functional/performance/security/compatibility",
            "priority": "P0/P1/P2/P3/P4"
        }
    ],
    "confidence_score": 0.95
}

注意：
- 仔细分析文档内容，准确识别需求和测试点
- 优先识别明确的功能需求和业务流程
- 为每个测试场景分配合适的测试类型和优先级
- 确保返回有效的JSON格式,去掉 ```json 和 ```
"""

    def _build_document_analysis_prompt(self, content: str, message: DocumentParseRequest) -> str:
        """构建文档分析提示"""
        return f"""
请分析以下需求文档，提取测试相关信息：

文档名称: {message.file_name}
文档类型: {message.document_type or "未知"}
分析目标: {message.analysis_target or "生成测试用例"}

文档内容：
{content[:10000]}  # 限制内容长度避免token超限

请根据文档内容，识别所有可测试的功能点和业务流程，生成对应的测试场景。
"""

    async def _run_ai_analysis(self, agent, prompt: str) -> str:
        """执行AI分析"""
        try:
            stream = agent.run_stream(task=prompt)
            async for event in stream:  # type: ignore
                # 流式消息，只是为了在前端界面流式显示
                if isinstance(event, ModelClientStreamingChunkEvent):
                    # 临时注释，不在前端显示流式内容
                    # await self.send_response(content=event.content, source=self.id.key)
                    continue

                # 最终的完整结果
                if isinstance(event, TaskResult):
                    messages = event.messages
                    # 从最后一条消息中获取完整内容
                    if messages and hasattr(messages[-1], 'content'):
                        return messages[-1].content

            # 如果没有获取到结果，返回默认值
            return """
                {
                    "document_type": "功能需求文档",
                    "title": "系统需求规格说明书",
                    "sections": [],
                    "requirements": [],
                    "test_scenarios": [],
                    "confidence_score": 0.8
                }
                """
        except Exception as e:
            logger.error(f"AI分析执行失败: {str(e)}")
            # 返回默认结果而不是抛出异常
            return """
{
    "document_type": "功能需求文档",
    "title": "解析失败",
    "sections": [],
    "requirements": [],
    "test_scenarios": [],
    "confidence_score": 0.5
}
"""

    def _parse_ai_analysis_result(self, ai_result: str, original_content: str) -> DocumentParseResult:
        """解析AI分析结果"""
        try:
            # 尝试解析JSON
            result_data = json.loads(ai_result.replace("```json", "").replace("```", ""))
            
            return DocumentParseResult(
                document_type=result_data.get("document_type", "unknown"),
                title=result_data.get("title", "未知文档"),
                content=original_content,
                sections=result_data.get("sections", []),
                requirements=result_data.get("requirements", []),
                test_scenarios=result_data.get("test_scenarios", []),
                confidence_score=result_data.get("confidence_score", 0.5)
            )
            
        except json.JSONDecodeError:
            logger.warning("AI返回结果不是有效JSON，使用默认解析")
            return DocumentParseResult(
                document_type="unknown",
                title="解析失败",
                content=original_content,
                sections=[],
                requirements=[],
                test_scenarios=[],
                confidence_score=0.3
            )

    async def _generate_test_cases_from_document(self, parse_result: DocumentParseResult, message: DocumentParseRequest) -> List[TestCaseData]:
        """从文档解析结果生成测试用例"""
        test_cases = []
        
        try:
            # 基于测试场景生成测试用例
            for scenario in parse_result.test_scenarios:
                test_case = TestCaseData(
                    title=scenario.get("title", "未命名测试用例"),
                    description=scenario.get("description", ""),
                    test_type=self._map_test_type(scenario.get("test_type", "functional")),
                    test_level=TestLevel.SYSTEM,
                    priority=self._map_priority(scenario.get("priority", "P2")),
                    input_source=InputSource.DOCUMENT,
                    ai_confidence=parse_result.confidence_score
                )
                test_cases.append(test_case)
            
            # 如果没有明确的测试场景，基于需求生成
            if not test_cases and parse_result.requirements:
                for req in parse_result.requirements:
                    test_case = TestCaseData(
                        title=f"测试需求: {req.get('title', '未命名需求')}",
                        description=req.get("description", ""),
                        test_type=TestType.FUNCTIONAL,
                        test_level=TestLevel.SYSTEM,
                        priority=self._map_priority(req.get("priority", "medium")),
                        input_source=InputSource.DOCUMENT,
                        source_metadata={
                            "document_name": message.file_name,
                            "requirement_id": req.get("id")
                        },
                        ai_confidence=parse_result.confidence_score
                    )
                    test_cases.append(test_case)
            
            logger.info(f"从文档生成了 {len(test_cases)} 个测试用例")
            return test_cases
            
        except Exception as e:
            logger.error(f"生成测试用例失败: {str(e)}")
            return []

    def _map_test_type(self, test_type_str: str) -> TestType:
        """映射测试类型"""
        mapping = {
            "functional": TestType.FUNCTIONAL,
            "performance": TestType.PERFORMANCE,
            "security": TestType.SECURITY,
            "compatibility": TestType.COMPATIBILITY,
            "usability": TestType.USABILITY,
            "interface": TestType.INTERFACE,
            "database": TestType.DATABASE
        }
        return mapping.get(test_type_str.lower(), TestType.FUNCTIONAL)

    def _map_priority(self, priority_str: str) -> Priority:
        """映射优先级"""
        mapping = {
            "high": Priority.P1,
            "medium": Priority.P2,
            "low": Priority.P3,
            "P0": Priority.P0,
            "P1": Priority.P1,
            "P2": Priority.P2,
            "P3": Priority.P3,
            "P4": Priority.P4
        }
        return mapping.get(priority_str, Priority.P2)

    async def _send_to_test_point_extractor(self, response: DocumentParseResponse):
        """发送到测试点提取智能体"""
        try:
            from app.core.messages.test_case import TestPointExtractionRequest

            # 构建需求解析结果
            requirement_analysis_result = {
                "source_type": "document",
                "document_name": response.file_name,
                "document_content": response.parse_result,
                "requirements": [tc.model_dump() for tc in response.test_cases],
                "business_processes": response.parse_result.get("business_processes", []),
                "functional_requirements": response.parse_result.get("functional_requirements", []),
                "non_functional_requirements": response.parse_result.get("non_functional_requirements", []),
                "constraints": response.parse_result.get("constraints", []),
                "assumptions": response.parse_result.get("assumptions", [])
            }

            extraction_request = TestPointExtractionRequest(
                session_id=response.session_id,
                requirement_analysis_result=requirement_analysis_result,
                extraction_config={
                    "enable_functional_testing": True,
                    "enable_non_functional_testing": True,
                    "enable_integration_testing": True,
                    "enable_acceptance_testing": True,
                    "enable_boundary_testing": True,
                    "enable_exception_testing": True,
                    "test_depth": "comprehensive"
                },
                test_strategy="document_driven"
            )

            await self.publish_message(
                extraction_request,
                topic_id=TopicId(type=TopicTypes.TEST_POINT_EXTRACTOR.value, source=self.id.key)
            )

            logger.info(f"已发送到测试点提取智能体: {response.session_id}")

        except Exception as e:
            logger.error(f"发送到测试点提取智能体失败: {str(e)}")

    async def _save_requirements_to_database(self, parse_result: DocumentParseResult, message: DocumentParseRequest) -> None:
        """保存需求到数据库"""
        try:
            await self.send_response(
                f"💾 开始保存 {len(parse_result.requirements)} 个需求到数据库...",
                region="process"
            )

            # 构建需求保存请求
            requirement_save_request = RequirementSaveRequest(
                session_id=message.session_id,
                document_id=str(uuid.uuid4()),
                file_name=message.file_name,
                file_path=message.file_path,
                requirements=parse_result.requirements,
                project_id=None,  # 使用默认项目
                ai_model_info={
                    "model": "deepseek-chat",
                    "generation_time": datetime.now().isoformat(),
                    "agent_version": "1.0",
                    "session_id": message.session_id,
                    "confidence_score": parse_result.confidence_score
                }
            )

            # 发送到需求存储智能体
            await self.publish_message(
                requirement_save_request,
                topic_id=TopicId(type=TopicTypes.REQUIREMENT_SAVER.value, source=self.id.key)
            )

            logger.info(f"已发送需求保存请求到需求存储智能体: {message.session_id}")

        except Exception as e:
            logger.error(f"保存需求到数据库失败: {str(e)}")
            await self.send_response(
                f"⚠️ 需求保存失败: {str(e)}",
                region="warning"
            )
