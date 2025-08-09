"""
导出功能API端点 (最终版)
"""
import uuid
from typing import List, Optional, Dict, Any
from fastapi import APIRouter, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field
from sqlalchemy import select, desc
from loguru import logger

from app.database.connection import db_manager
from app.database.models.test_case import ExportRecord, Project, TestCase, ExportType, ExportStatus

router = APIRouter()

class ExportRequest(BaseModel):
    """导出请求"""
    export_type: ExportType = Field(..., description="导出类型")
    test_case_ids: List[str] = Field(..., description="测试用例ID列表")
    project_id: Optional[str] = Field(None, description="项目ID")
    session_id: Optional[str] = Field(None, description="会话ID")
    export_config: Optional[Dict[str, Any]] = Field(None, description="导出配置")

class ExportResponse(BaseModel):
    """导出响应"""
    id: str
    export_type: ExportType
    test_case_count: int
    project_id: Optional[str]
    project_name: Optional[str] = None
    session_id: Optional[str]
    file_name: str
    file_path: str
    file_size: Optional[int]
    status: ExportStatus
    created_at: str

    class Config:
        from_attributes = True

class ExportListResponse(BaseModel):
    """导出列表响应"""
    items: List[ExportResponse]
    total: int

@router.post("/", response_model=ExportResponse)
async def create_export(request: ExportRequest, background_tasks: BackgroundTasks):
    """创建导出任务"""
    try:
        async with db_manager.get_session() as session:
            # 验证测试用例存在
            test_cases_result = await session.execute(
                select(TestCase).where(TestCase.id.in_(request.test_case_ids))
            )
            test_cases = test_cases_result.scalars().all()
            
            if len(test_cases) != len(request.test_case_ids):
                raise HTTPException(status_code=404, detail="部分测试用例不存在")
            
            # 验证项目存在（如果提供）
            project_name = None
            if request.project_id:
                project_result = await session.execute(
                    select(Project).where(Project.id == request.project_id)
                )
                project = project_result.scalar_one_or_none()
                if not project:
                    raise HTTPException(status_code=404, detail="项目不存在")
                project_name = project.name
            
            # 生成文件名
            import datetime
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            file_extension = {
                ExportType.EXCEL: "xlsx",
                ExportType.WORD: "docx",
                ExportType.PDF: "pdf"
            }[request.export_type]
            
            file_name = f"test_cases_{timestamp}.{file_extension}"
            file_path = f"/exports/{file_name}"
            
            # 创建导出记录
            export_record = ExportRecord(
                id=str(uuid.uuid4()),
                export_type=request.export_type,
                test_case_ids=request.test_case_ids,
                project_id=request.project_id,
                session_id=request.session_id,
                file_name=file_name,
                file_path=file_path,
                export_config=request.export_config,
                status=ExportStatus.PENDING
            )
            
            session.add(export_record)
            await session.commit()
            await session.refresh(export_record)
            
            # 启动后台导出任务
            background_tasks.add_task(
                _process_export_task,
                export_record.id,
                request.export_type,
                test_cases,
                file_path,
                request.export_config
            )
            
            return ExportResponse(
                id=export_record.id,
                export_type=export_record.export_type,
                test_case_count=len(request.test_case_ids),
                project_id=export_record.project_id,
                project_name=project_name,
                session_id=export_record.session_id,
                file_name=export_record.file_name,
                file_path=export_record.file_path,
                file_size=export_record.file_size,
                status=export_record.status,
                created_at=export_record.created_at.isoformat()
            )
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"创建导出任务失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"创建导出任务失败: {str(e)}")

@router.get("/", response_model=ExportListResponse)
async def get_exports(
    project_id: Optional[str] = None,
    session_id: Optional[str] = None,
    export_type: Optional[ExportType] = None,
    status: Optional[ExportStatus] = None
):
    """获取导出记录列表"""
    try:
        async with db_manager.get_session() as session:
            query = select(ExportRecord)
            
            if project_id:
                query = query.where(ExportRecord.project_id == project_id)
            
            if session_id:
                query = query.where(ExportRecord.session_id == session_id)
            
            if export_type:
                query = query.where(ExportRecord.export_type == export_type)
            
            if status:
                query = query.where(ExportRecord.status == status)
            
            query = query.order_by(desc(ExportRecord.created_at))
            
            result = await session.execute(query)
            exports = result.scalars().all()
            
            # 构建响应
            items = []
            for export_record in exports:
                # 获取项目名称
                project_name = None
                if export_record.project_id:
                    project_result = await session.execute(
                        select(Project.name).where(Project.id == export_record.project_id)
                    )
                    project_name = project_result.scalar()
                
                test_case_count = len(export_record.test_case_ids) if export_record.test_case_ids else 0
                
                items.append(ExportResponse(
                    id=export_record.id,
                    export_type=export_record.export_type,
                    test_case_count=test_case_count,
                    project_id=export_record.project_id,
                    project_name=project_name,
                    session_id=export_record.session_id,
                    file_name=export_record.file_name,
                    file_path=export_record.file_path,
                    file_size=export_record.file_size,
                    status=export_record.status,
                    created_at=export_record.created_at.isoformat()
                ))
            
            return ExportListResponse(
                items=items,
                total=len(items)
            )
            
    except Exception as e:
        logger.error(f"获取导出列表失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"获取导出列表失败: {str(e)}")

@router.get("/{export_id}", response_model=ExportResponse)
async def get_export(export_id: str):
    """获取导出记录详情"""
    try:
        async with db_manager.get_session() as session:
            result = await session.execute(
                select(ExportRecord).where(ExportRecord.id == export_id)
            )
            export_record = result.scalar_one_or_none()
            
            if not export_record:
                raise HTTPException(status_code=404, detail="导出记录不存在")
            
            # 获取项目名称
            project_name = None
            if export_record.project_id:
                project_result = await session.execute(
                    select(Project.name).where(Project.id == export_record.project_id)
                )
                project_name = project_result.scalar()
            
            test_case_count = len(export_record.test_case_ids) if export_record.test_case_ids else 0
            
            return ExportResponse(
                id=export_record.id,
                export_type=export_record.export_type,
                test_case_count=test_case_count,
                project_id=export_record.project_id,
                project_name=project_name,
                session_id=export_record.session_id,
                file_name=export_record.file_name,
                file_path=export_record.file_path,
                file_size=export_record.file_size,
                status=export_record.status,
                created_at=export_record.created_at.isoformat()
            )
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取导出详情失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"获取导出详情失败: {str(e)}")

@router.get("/{export_id}/download")
async def download_export(export_id: str):
    """下载导出文件"""
    try:
        async with db_manager.get_session() as session:
            result = await session.execute(
                select(ExportRecord).where(ExportRecord.id == export_id)
            )
            export_record = result.scalar_one_or_none()
            
            if not export_record:
                raise HTTPException(status_code=404, detail="导出记录不存在")
            
            if export_record.status != ExportStatus.COMPLETED:
                raise HTTPException(status_code=400, detail="导出文件尚未完成")
            
            import os
            full_path = os.path.join(os.getcwd(), export_record.file_path.lstrip('/'))
            
            if not os.path.exists(full_path):
                raise HTTPException(status_code=404, detail="导出文件不存在")
            
            return FileResponse(
                path=full_path,
                filename=export_record.file_name,
                media_type='application/octet-stream'
            )
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"下载导出文件失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"下载导出文件失败: {str(e)}")

@router.delete("/{export_id}")
async def delete_export(export_id: str):
    """删除导出记录"""
    try:
        async with db_manager.get_session() as session:
            result = await session.execute(
                select(ExportRecord).where(ExportRecord.id == export_id)
            )
            export_record = result.scalar_one_or_none()
            
            if not export_record:
                raise HTTPException(status_code=404, detail="导出记录不存在")
            
            # 删除文件
            import os
            if export_record.file_path:
                full_path = os.path.join(os.getcwd(), export_record.file_path.lstrip('/'))
                if os.path.exists(full_path):
                    os.remove(full_path)
            
            # 删除记录
            await session.delete(export_record)
            await session.commit()
            
            return {"message": "导出记录删除成功"}
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"删除导出记录失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"删除导出记录失败: {str(e)}")

async def _process_export_task(
    export_id: str,
    export_type: ExportType,
    test_cases: List[TestCase],
    file_path: str,
    export_config: Optional[Dict[str, Any]]
):
    """后台处理导出任务"""
    try:
        async with db_manager.get_session() as session:
            # 更新状态为处理中
            export_result = await session.execute(
                select(ExportRecord).where(ExportRecord.id == export_id)
            )
            export_record = export_result.scalar_one()
            export_record.status = ExportStatus.PROCESSING
            await session.commit()
            
            # 创建导出目录
            import os
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            full_path = os.path.join(os.getcwd(), file_path.lstrip('/'))
            
            # 根据导出类型处理
            if export_type == ExportType.EXCEL:
                await _export_to_excel(test_cases, full_path, export_config)
            elif export_type == ExportType.WORD:
                await _export_to_word(test_cases, full_path, export_config)
            elif export_type == ExportType.PDF:
                await _export_to_pdf(test_cases, full_path, export_config)
            
            # 获取文件大小
            file_size = os.path.getsize(full_path)
            
            # 更新状态为完成
            export_record.status = ExportStatus.COMPLETED
            export_record.file_size = file_size
            await session.commit()
            
    except Exception as e:
        logger.error(f"导出任务处理失败: {str(e)}")
        # 更新状态为失败
        async with db_manager.get_session() as session:
            export_result = await session.execute(
                select(ExportRecord).where(ExportRecord.id == export_id)
            )
            export_record = export_result.scalar_one_or_none()
            if export_record:
                export_record.status = ExportStatus.FAILED
                await session.commit()

async def _export_to_excel(test_cases: List[TestCase], file_path: str, config: Optional[Dict[str, Any]]):
    """导出到Excel"""
    try:
        import pandas as pd
        
        # 准备数据
        data = []
        for tc in test_cases:
            data.append({
                '测试用例ID': tc.id,
                '标题': tc.title,
                '描述': tc.description or '',
                '前置条件': tc.preconditions or '',
                '测试步骤': str(tc.test_steps) if tc.test_steps else '',
                '预期结果': tc.expected_results or '',
                '测试类型': tc.test_type,
                '测试级别': tc.test_level,
                '优先级': tc.priority,
                '状态': tc.status,
                'AI生成': '是' if tc.ai_generated else '否',
                'AI置信度': tc.ai_confidence or '',
                '创建时间': tc.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                '更新时间': tc.updated_at.strftime('%Y-%m-%d %H:%M:%S')
            })
        
        # 创建DataFrame并导出
        df = pd.DataFrame(data)
        df.to_excel(file_path, index=False, engine='openpyxl')
        
    except Exception as e:
        logger.error(f"Excel导出失败: {str(e)}")
        raise

async def _export_to_word(test_cases: List[TestCase], file_path: str, config: Optional[Dict[str, Any]]):
    """导出到Word"""
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('测试用例报告', 0)
        
        for i, tc in enumerate(test_cases, 1):
            doc.add_heading(f'{i}. {tc.title}', level=1)
            
            # 基本信息表格
            table = doc.add_table(rows=8, cols=2)
            table.style = 'Table Grid'
            
            cells = table.rows[0].cells
            cells[0].text = '测试用例ID'
            cells[1].text = tc.id
            
            cells = table.rows[1].cells
            cells[0].text = '测试类型'
            cells[1].text = tc.test_type
            
            cells = table.rows[2].cells
            cells[0].text = '测试级别'
            cells[1].text = tc.test_level
            
            cells = table.rows[3].cells
            cells[0].text = '优先级'
            cells[1].text = tc.priority
            
            cells = table.rows[4].cells
            cells[0].text = '状态'
            cells[1].text = tc.status
            
            cells = table.rows[5].cells
            cells[0].text = '前置条件'
            cells[1].text = tc.preconditions or ''
            
            cells = table.rows[6].cells
            cells[0].text = '测试步骤'
            cells[1].text = str(tc.test_steps) if tc.test_steps else ''
            
            cells = table.rows[7].cells
            cells[0].text = '预期结果'
            cells[1].text = tc.expected_results or ''
            
            doc.add_paragraph()
        
        doc.save(file_path)
        
    except Exception as e:
        logger.error(f"Word导出失败: {str(e)}")
        raise

async def _export_to_pdf(test_cases: List[TestCase], file_path: str, config: Optional[Dict[str, Any]]):
    """导出到PDF"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        
        doc = SimpleDocTemplate(file_path, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # 标题
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # 居中
        )
        story.append(Paragraph('测试用例报告', title_style))
        story.append(Spacer(1, 20))
        
        # 测试用例内容
        for i, tc in enumerate(test_cases, 1):
            # 用例标题
            story.append(Paragraph(f'{i}. {tc.title}', styles['Heading2']))
            
            # 用例信息表格
            data = [
                ['测试用例ID', tc.id],
                ['测试类型', tc.test_type],
                ['测试级别', tc.test_level],
                ['优先级', tc.priority],
                ['状态', tc.status],
                ['前置条件', tc.preconditions or ''],
                ['测试步骤', str(tc.test_steps) if tc.test_steps else ''],
                ['预期结果', tc.expected_results or '']
            ]
            
            table = Table(data, colWidths=[2*inch, 4*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.grey),
                ('TEXTCOLOR', (0, 0), (0, -1), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('BACKGROUND', (1, 0), (1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(table)
            story.append(Spacer(1, 20))
        
        doc.build(story)
        
    except Exception as e:
        logger.error(f"PDF导出失败: {str(e)}")
        raise
